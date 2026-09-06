from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.image.image import Image
from docx.shared import Pt
from docx.oxml.ns import qn


WORKSPACE = Path(__file__).resolve().parents[1]
DOCX_PATH = next(
    p for p in (WORKSPACE / "论文").glob("*上下文工程方法研究.docx")
    if not p.name.startswith("~$")
)
FIGURES = {
    "!图3-1": ("第3章", "图3-1-方法总体框架.png", "图3-1 方法总体框架"),
    "!图3-2": ("第3章", "图3-2-六阶段流程约束架构.png", "图3-2 六阶段流程约束架构"),
    "!图3-3": ("第3章", "图3-3-能力代码双层本体模型.png", "图3-3 能力-代码双层本体模型"),
    "!图3-4": ("第3章", "图3-4-business_map生成流程.png", "图3-4 business_map 生成流程"),
    "!图3-5": ("第3章", "图3-5-图谱驱动任务上下文生成流水线.png", "图3-5 图谱驱动的任务上下文生成流水线"),
    "!图4-1": ("第4章", "图4-1-wpw系统三层架构.png", "图4-1 wpw 系统三层架构"),
    "!图4-2": ("第4章", "图4-2-工作流引擎实现架构.png", "图4-2 工作流引擎实现架构"),
    "!图4-3": ("第4章", "图4-3-阶段制品生成与确认流程.png", "图4-3 阶段制品生成与确认流程"),
    "!图4-4": ("第4章", "图4-4-DAG驱动的阶段门禁.png", "图4-4 DAG 驱动的阶段门禁"),
    "!图4-5": ("第4章", "图4-5-加权双向BFS子图扩展示意.png", "图4-5 加权双向 BFS 子图扩展示意"),
    "!图5-1": ("第5章美食推荐系统", "图5-1-系统功能结构图.png", "图5-1 系统功能结构图"),
    "!图5-2": ("第5章美食推荐系统", "图5-2-管理员用例图.png", "图5-2 管理员用例图"),
    "!图5-3": ("第5章美食推荐系统", "图5-3-用户用例图.png", "图5-3 用户用例图"),
    "!图5-4": ("第5章美食推荐系统", "图5-4-管理员信息实体图.png", "图5-4 管理员信息实体图"),
    "!图5-5": ("第5章美食推荐系统", "图5-5-用户信息实体图.png", "图5-5 用户信息实体图"),
    "!图5-6": ("第5章美食推荐系统", "图5-6-美食信息实体图.png", "图5-6 美食信息实体图"),
    "!图5-7": ("第5章美食推荐系统", "图5-7-订单信息实体图.png", "图5-7 订单信息实体图"),
    "!图5-8": ("第5章美食推荐系统", "图5-8-用户注册界面.png", "图5-8 用户注册界面"),
    "!图5-9": ("第5章美食推荐系统", "图5-9-用户登录界面.png", "图5-9 用户登录界面"),
    "!图5-10": ("第5章美食推荐系统", "图5-10-小程序首页界面.png", "图5-10 小程序首页界面"),
    "!图5-11": ("第5章美食推荐系统", "图5-11-美食信息界面.png", "图5-11 美食信息界面"),
    "!图5-12": ("第5章美食推荐系统", "图5-12-购物车界面.png", "图5-12 购物车界面"),
    "!图5-13": ("第5章美食推荐系统", "图5-13-我的功能界面.png", "图5-13 我的功能界面"),
    "!图5-14": ("第5章美食推荐系统", "图5-14-管理员登录界面.png", "图5-14 管理员登录界面"),
    "!图5-15": ("第5章美食推荐系统", "图5-15-管理员功能界面.png", "图5-15 管理员功能界面"),
    "!图5-16": ("第5章美食推荐系统", "图5-16-用户管理界面.png", "图5-16 用户管理界面"),
    "!图5-17": ("第5章美食推荐系统", "图5-17-美食分类管理界面.png", "图5-17 美食分类管理界面"),
    "!图5-18": ("第5章美食推荐系统", "图5-18-美食信息管理界面.png", "图5-18 美食信息管理界面"),
    "!图5-19": ("第5章美食推荐系统", "图5-19-系统管理界面.png", "图5-19 系统管理界面"),
    "!图5-20": ("第5章美食推荐系统", "图5-20-订单管理界面.png", "图5-20 订单管理界面"),
}


def insert_after(anchor, paragraph):
    anchor._p.addnext(paragraph._p)


document = Document(str(DOCX_PATH))
max_text_width = min(
    section.page_width - section.left_margin - section.right_margin
    for section in document.sections
)
placeholders = [p for p in document.paragraphs if p.text.strip() in FIGURES]
if not placeholders:
    print("No Chapter 3 figure placeholders found; document may already be synchronized.")
    raise SystemExit(0)

for placeholder in placeholders:
    key = placeholder.text.strip()
    figure_dir, image_name, caption_text = FIGURES[key]
    image_path = WORKSPACE / "论文" / "图表" / figure_dir / image_name
    if not image_path.exists():
        raise FileNotFoundError(image_path)

    figure = document.add_paragraph()
    figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure.paragraph_format.space_before = 0
    figure.paragraph_format.space_after = 0
    figure.paragraph_format.keep_with_next = True
    native_width = Image.from_file(str(image_path)).width
    figure.add_run().add_picture(str(image_path), width=min(native_width, max_text_width))

    caption = document.add_paragraph(style="u图标题")
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_before = Pt(1.2)
    caption.paragraph_format.space_after = Pt(12)
    caption.paragraph_format.line_spacing = 1.0
    caption.add_run(caption_text)
    for run in caption.runs:
        run.font.name = "Times New Roman"
        run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "黑体")
        run.font.size = Pt(10.5)
        run.bold = True

    anchor = placeholder._p
    parent = anchor.getparent()
    insert_after(placeholder, figure)
    insert_after(figure, caption)
    parent.remove(anchor)

document.save(str(DOCX_PATH))
print(f"Embedded {len(placeholders)} chapter figures in {DOCX_PATH}")
