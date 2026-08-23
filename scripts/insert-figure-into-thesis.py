from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement


WORKSPACE = Path(r"F:\project\web-project-workflow")
DOCX_PATH = WORKSPACE / r"论文\基于大模型的美食推荐系统设计与实现--面向AI辅助开发的上下文工程方法研究.docx"
OUTPUT_PATH = WORKSPACE / r"论文\基于大模型的美食推荐系统设计与实现--面向AI辅助开发的上下文工程方法研究-已插入图3-1.docx"
IMAGE_PATH = WORKSPACE / r"论文\图表\图3-1-方法总体框架.png"


def insert_after(paragraph, new_paragraph):
    paragraph._p.addnext(new_paragraph._p)


def set_keep_with_next(paragraph, enabled=True):
    ppr = paragraph._p.get_or_add_pPr()
    tag = ppr.find(qn("w:keepNext"))
    if enabled and tag is None:
        ppr.append(OxmlElement("w:keepNext"))
    elif not enabled and tag is not None:
        ppr.remove(tag)


def qn(tag):
    from docx.oxml.ns import qn as _qn
    return _qn(tag)


document = Document(str(DOCX_PATH))
placeholder = next(
    (p for p in document.paragraphs if p.text.strip().startswith("图3-1")),
    None,
)
if placeholder is None:
    raise RuntimeError("未找到图3-1占位段")

previous = placeholder._p.getprevious()
parent = placeholder._p.getparent()
parent.remove(placeholder._p)

figure = document.add_paragraph()
figure.alignment = WD_ALIGN_PARAGRAPH.CENTER
figure.paragraph_format.space_before = Pt(6)
figure.paragraph_format.space_after = Pt(2)
figure.paragraph_format.keep_with_next = True
run = figure.add_run()
run.add_picture(str(IMAGE_PATH), width=Inches(6.2))

caption = document.add_paragraph()
caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
caption.paragraph_format.space_before = Pt(0)
caption.paragraph_format.space_after = Pt(8)
caption.paragraph_format.keep_with_next = True
caption_run = caption.add_run("图3-1  面向 AI 辅助开发的上下文工程方法总体框架")
caption_run.font.name = "宋体"
caption_run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
caption_run.font.size = Pt(10.5)

if previous is None:
    body = document._body._element
    body.insert(0, figure._p)
    body.insert(1, caption._p)
else:
    previous.addnext(figure._p)
    figure._p.addnext(caption._p)

document.save(str(OUTPUT_PATH))
print(f"Inserted figure into {OUTPUT_PATH}")
